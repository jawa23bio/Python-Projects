import os
import time
from flask import Flask, request, render_template_string, send_file, session, redirect, url_for
from docx import Document
import requests
import logging
import tiktoken
from io import BytesIO
from dotenv import load_dotenv

# === Load environment variables (API keys, secrets, etc.) ===
load_dotenv()

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev")  # Use strong secret in prod!

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# === Azure OpenAI Configuration ===
API_KEY = os.environ.get("AZURE_OPENAI_API_KEY")
DEPLOYMENT_ID = os.environ.get("AZURE_DEPLOYMENT_ID", "gpt-4")
API_VERSION = os.environ.get("AZURE_API_VERSION", "2025-01-01-preview")
BASE_URL = os.environ.get("AZURE_BASE_URL")
if not API_KEY or not BASE_URL:
    raise ValueError("API_KEY and BASE_URL must be set as environment variables.")

COMPLETION_URL = f"{BASE_URL}/deployments/{DEPLOYMENT_ID}/chat/completions?api-version={API_VERSION}"

# === Prompt Templates (same as before, short version here) ===
PURPOSE_PROMPT = """
You are a technical documentation assistant. Based on all available documentation, write a concise, human-readable summary of this dashboard’s primary purpose, intended audience, data sources, and unique features. Do not use a fill-in-the-blank template. Be specific and contextual; mention what makes this dashboard unique or important for its intended users.

--- Dashboard Summaries ---
{aggregated_summary}
--- End Summaries ---
"""

FILTERS_PROMPT_TEMPLATE = """
You are a documentation assistant. Given the following dashboard documentation (including all extracted visualization details and any other available sections), identify the filtering scheme applied for the visualizations.
Check the Used Filter field for all the visualizations and mention the name of the filtering scheme that was used. If the same filtering scheme is present for all the visualization, say it was global filtering scheme present for all visualizations in the analysis.
Only list columns that are actually referenced as filters or filtering columns in the provided documentation. Do NOT invent or generalize.
If you cannot find evidence of a global filtering scheme, mention the visualization-specific filters.

--- Dashboard Content Start ---
{dashboard_content}
--- Dashboard Content End ---
"""

# === Utility functions (unchanged except to fit this structure) ===

def get_dashboard_content_for_filters(viz_blocks, extra_sections=None):
    blocks = []
    if extra_sections:
        blocks.extend(extra_sections)
    for viz in viz_blocks:
        blocks.append(f"---\nTitle: {viz['title']}\nPage: {viz['subpage']}\nDetails:\n{viz['details'].strip()}")
    return "\n\n".join(blocks)

def generate_filters_section_with_gpt(viz_blocks, extra_sections=None):
    dashboard_content = get_dashboard_content_for_filters(viz_blocks, extra_sections)
    prompt = FILTERS_PROMPT_TEMPLATE.format(dashboard_content=dashboard_content)
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "messages": [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.2,
        "max_tokens": 700
    }
    response = requests.post(COMPLETION_URL, headers=headers, json=data)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"].strip()

def extract_scripts(file_stream, section_heading):
    doc = Document(file_stream)
    scripts = []
    current_script = None
    in_section = False
    capture_code = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name

        if style == "Heading 2":
            if section_heading.upper() in text.upper():
                in_section = True
                current_script = None
                capture_code = False
                continue
            elif in_section:
                break

        if not in_section:
            continue

        if style == "Heading 3":
            if current_script:
                scripts.append(current_script)
            current_script = {"name": text, "description": "", "code": ""}
            capture_code = False
            continue

        if not current_script:
            continue

        if text.lower() == "description":
            capture_code = False
            continue

        if text.lower() == "script parameters":
            capture_code = False
            continue

        if text.lower() == "script definition":
            capture_code = True
            continue

        if not capture_code and not current_script["description"]:
            current_script["description"] = text
            continue

        if capture_code:
            current_script["code"] += text + "\n"

    if current_script:
        scripts.append(current_script)
    return scripts

def extract_from_template_overview(file_stream, max_chars=300000):
    doc = Document(file_stream)
    capture = False
    content_lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if 'Template Overview' in text:
            capture = True
            continue
        if capture:
            content_lines.append(text)
            if sum(len(line) for line in content_lines) > max_chars:
                break
    return "\n".join(content_lines)

def extract_document_properties_table(file_stream):
    doc = Document(file_stream)
    capture = False
    properties = []

    for para in doc.paragraphs:
        if para.text.strip().lower() == "document properties":
            capture = True
            break

    if not capture:
        return []

    for table in doc.tables:
        if table.cell(0, 0).text.strip().lower() == "property name":
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if cells:
                    properties.append(cells)
            break
    return properties

def chunk_by_tokens(text, model_name="gpt-4"):
    safe_max = {
        "gpt-4-32k": 24000,
        "gpt-4-1106-preview": 24000,
        "gpt-4-0125-preview": 24000,
        "gpt-4": 12000,
        "gpt-3.5-turbo-16k": 12000,
        "gpt-3.5-turbo": 7000,
    }
    max_tokens = safe_max.get(model_name, 7000)
    enc = tiktoken.encoding_for_model(model_name if model_name in safe_max else "gpt-4")
    tokens = enc.encode(text)
    chunks = []
    start = 0
    while start < len(tokens):
        end = min(start + max_tokens, len(tokens))
        chunk = enc.decode(tokens[start:end])
        chunks.append(chunk)
        start = end
    return chunks

def summarize_chunk_lightly(chunk, chunk_number):
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": f"Summarize this dashboard documentation chunk (Chunk {chunk_number}) focusing on: dashboard purpose, data sources, key elements, and document properties.\n\n{chunk}"}
    ]
    data = {
        "messages": messages,
        "temperature": 0.3,
        "max_tokens": 1000
    }
    try:
        response = requests.post(COMPLETION_URL, headers=headers, json=data)
        response.raise_for_status()
        time.sleep(1.2)
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        logger.error(f"Chunk {chunk_number} summarization error: {e}")
        return f"❌ Error in chunk {chunk_number}: {str(e)}"

def generate_final_summary_from_chunks(aggregated_summary):
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    final_prompt = PURPOSE_PROMPT.format(aggregated_summary=aggregated_summary)
    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": final_prompt}
    ]
    data = {
        "messages": messages,
        "temperature": 0.2,
        "max_tokens": 600
    }
    try:
        response = requests.post(COMPLETION_URL, headers=headers, json=data)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        logger.error(f"Final summary generation error: {e}")
        return f"❌ Error in final summary: {str(e)}"

def generate_property_descriptions_with_gpt(raw_properties):
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    if not raw_properties:
        return "❌ No document properties table found."
    header = "Property Name | Type | Value | Script to Execute\n"
    separator = "------------- | ---- | ----- | ------------------\n"
    rows = "\n".join(" | ".join(cell.strip() for cell in row[:4]) for row in raw_properties if len(row) >= 4)
    markdown_table = header + separator + rows
    prompt = f"""
You are a documentation assistant. Below is a table of document properties from a dashboard configuration:

{markdown_table}

Your task is to write clean, structured descriptions for each property in the following format — do NOT include any markdown tables:

- **<Property Name>**
  - **Purpose**: Briefly describe what it controls in the dashboard.
  - **Usage**: Where and how it is used (e.g., visualizations, filters, scripts).
  - **Linked Script**: Script name from the 'Script to Execute' column. If none, say "None".

Only output descriptions. Do not repeat or include the table again.
"""
    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": prompt}
    ]
    data = {
        "messages": messages,
        "temperature": 0.2,
        "max_tokens": 1500
    }
    try:
        response = requests.post(COMPLETION_URL, headers=headers, json=data)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        logger.error(f"Document property description error: {e}")
        return f"❌ Error generating document property descriptions: {str(e)}"

def generate_script_descriptions(scripts, script_type="IronPython"):
    if not scripts:
        return f"❌ No {script_type} scripts found."
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    descriptions = []
    for script in scripts:
        prompt = (
            f"You are a documentation assistant.\n"
            f"Below is a {script_type} script used in a dashboard.\n\n"
            f"Script Name: {script['name']}\n"
            f"Description: {script.get('description', '').strip()}\n"
            f"Code:\n{script['code']}\n\n"
            f"Please summarize what this script does and why it is necessary in 2-3 sentences. "
            f"Do NOT quote the code."
        )
        data = {
            "messages": [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.3,
            "max_tokens": 300
        }
        try:
            response = requests.post(COMPLETION_URL, headers=headers, json=data)
            response.raise_for_status()
            content = response.json()["choices"][0]["message"]["content"]
            descriptions.append(f"### {script['name']}\n{content}\n")
        except Exception as e:
            logger.error(f"Error summarizing {script_type} script '{script['name']}': {e}")
            descriptions.append(f"### {script['name']}\n❌ Error summarizing this script: {str(e)}\n")
    return "\n".join(descriptions)

def extract_visualizations_from_pages(file_stream):
    doc = Document(file_stream)
    visualizations = []
    in_pages_section = False
    current_page = ""
    current_subpage = ""
    current_viz = None
    collecting_details = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name
        if style == "Heading 1":
            if "pages" in text.lower():
                in_pages_section = True
                current_page = text
            else:
                in_pages_section = False
                current_page = ""
                current_subpage = ""
        elif not in_pages_section:
            continue
        elif style == "Heading 2":
            current_subpage = text
        elif style == "Heading 3":
            if current_viz:
                visualizations.append(current_viz)
            current_viz = {
                "title": text,
                "page": current_page,
                "subpage": current_subpage,
                "details": ""
            }
            collecting_details = True
        elif collecting_details and current_viz:
            if style.startswith("Heading"):
                collecting_details = False
            else:
                current_viz["details"] += text + "\n"
    if current_viz:
        visualizations.append(current_viz)
    return visualizations

def generate_visualization_descriptions_from_details(visualizations):
    if not visualizations:
        return "❌ No visualizations found."
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    summaries = []
    for viz in visualizations:
        prompt = f"""
You are a documentation assistant.

Below is the extracted detail of a visualization from a dashboard:

Title: {viz['title']}
Page: {viz['subpage']}
Details:
{viz['details']}

Based on the above, generate the following structured summary:

o Purpose: {{Provide a summarized description of all information listed for the respective visualization}}
o Location: {{Use Pages column for the respective visualization from the Pages/Tabs table}}
o Type of visualization: {{Lookup Visualization Type for the respective visualization.}}
o Data source: {{Look up the Used Data Table or Data Table field. Include all the relevant data tables. You can also lookup Pages/Tabs table for the respective visualization.}}
o Include Limit by: {{Lookup 'Data Limit (marking)' for the respective visualization. In the case of Text area visualizations, check in all the fields starting with 'Limit data'}}
o Markings: {{Lookup Used Marking for the respective visualization.}}
o Filtering schemes: {{Lookup Used Filter for the respective visualization.}}

Only fill values based on actual content. Do not assume. If a field is blank or 'None', say None.
"""
        data = {
            "messages": [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.2,
            "max_tokens": 600
        }
        try:
            response = requests.post(COMPLETION_URL, headers=headers, json=data)
            response.raise_for_status()
            summaries.append(f"### {viz['title']} (Page: {viz['subpage']})\n{response.json()['choices'][0]['message']['content']}\n")
        except Exception as e:
            logger.error(f"Error summarizing visualization '{viz['title']}': {e}")
            summaries.append(f"### {viz['title']} (Page: {viz['subpage']})\n❌ Error: {str(e)}\n")
    if len(visualizations) > 5:
        summaries.append("[Please check if all visualizations are included]")
    return "\n".join(summaries)

def extract_data_table_blocks(file_stream):
    doc = Document(file_stream)
    blocks = []
    in_data_tables = False
    current_table = None
    buffer = []
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name
        if style == "Heading 1" and "data tables" in text.lower():
            in_data_tables = True
            continue
        if not in_data_tables:
            continue
        if style == "Heading 2":
            if current_table:
                blocks.append({"title": current_table, "raw": "\n".join(buffer).strip()})
            current_table = text
            buffer = []
            continue
        if style == "Heading 1":
            if current_table:
                blocks.append({"title": current_table, "raw": "\n".join(buffer).strip()})
            break
        buffer.append(text)
    return blocks

def gpt_summarize_table(table_block):
    prompt = f"""
You are a documentation assistant. Given the following raw data table configuration text, extract and summarize the content into this markdown template:

---
### {table_block['title']}
{{One/two sentence human summary of this table's purpose.}}

**Data Table Source (Connection Path):**
- (one bullet for each unique data source as it is mentioned in the documentation)

**Data Table Transformations:**
For each transformation in the table, list the following:
1. **Transformation Name – [Column Name(s) Affected]**
   - **Purpose:** [Purpose of this transformation, as specifically as possible.]
   - **Expression Used:**
       ```r
       [Code or formula exactly as it appears in the documentation.]
       ```

Do **not** summarize or abbreviate code. Do not omit code blocks. List every transformation step present, even if repetitive.

**Data Table Relations:**
- (One bullet per relation. Say N/A if none.)

**On-Demand Settings:**
- (List each if present. Say N/A if not.)
---

Raw content:
\"\"\"
{table_block['raw']}
\"\"\"

Only use information in the raw content. Never generalize or invent.
Structure your output exactly as shown above, repeating the structure for each transformation.
"""
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "messages": [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.2,
        "max_tokens": 1600
    }
    response = requests.post(COMPLETION_URL, headers=headers, json=data)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]

def summarize_all_tables(file_stream):
    file_stream.seek(0)
    blocks = extract_data_table_blocks(file_stream)
    summaries = []
    for block in blocks:
        try:
            summary = gpt_summarize_table(block)
            summaries.append(summary)
        except Exception as e:
            summaries.append(f"### {block['title']}\n❌ Error summarizing this table: {e}")
    return "\n\n".join(summaries)

# ==== FLASK ROUTES ==== #

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        if "doc" not in request.files:
            return "❌ No file uploaded."
        doc_file = request.files["doc"]
        if not doc_file.filename.endswith(".docx"):
            return "❌ Only .docx files are supported."
        doc_bytes = doc_file.read()
        doc_buffer = BytesIO(doc_bytes)
        doc_buffer.seek(0)

        overview_text = extract_from_template_overview(doc_buffer)
        doc_buffer.seek(0)
        raw_properties = extract_document_properties_table(doc_buffer)
        doc_buffer.seek(0)
        iron_scripts = extract_scripts(doc_buffer, "Iron Python Scripts")
        doc_buffer.seek(0)
        javascript_scripts = extract_scripts(doc_buffer, "JavaScripts")
        doc_buffer.seek(0)
        viz_blocks = extract_visualizations_from_pages(doc_buffer)
        doc_buffer.seek(0)
        filters_section = generate_filters_section_with_gpt(viz_blocks)

        data_table_gpt_summary = summarize_all_tables(doc_buffer)
        doc_buffer.seek(0)
        # Summarize overview
        token_chunks = chunk_by_tokens(overview_text, model_name="gpt-4")
        chunk_summaries = [summarize_chunk_lightly(chunk, i + 1) for i, chunk in enumerate(token_chunks)]
        aggregated = "\n\n".join(f"--- Chunk {i+1} ---\n{summary}" for i, summary in enumerate(chunk_summaries))
        final_summary = generate_final_summary_from_chunks(aggregated)
        formatted_properties = generate_property_descriptions_with_gpt(raw_properties)
        ironpython_summary = generate_script_descriptions(iron_scripts, "IronPython")
        javascript_summary = generate_script_descriptions(javascript_scripts, "JavaScript")
        viz_summary = generate_visualization_descriptions_from_details(viz_blocks)

        markdown_doc = (
            f"## 📄 Dashboard Purpose\n{final_summary}\n\n"
            f"## ⚙️ Document Properties\n{formatted_properties}\n\n"
            f"## 🐍 IronPython Scripts\n{ironpython_summary}\n\n"
            f"## 📜 JavaScripts\n{javascript_summary}\n\n"
            f"## 📊 Visualizations\n{viz_summary}\n\n"
            f"## 🗃️ Data Table Summaries (GPT-4)\n{data_table_gpt_summary}\n\n"
            f"## 🪄 Filters\n{filters_section}\n"
        )

        session["markdown_doc"] = markdown_doc

        # Render HTML preview
        return render_template_string("""
        <h2>📄 Dashboard Purpose</h2>
        <pre style="white-space: pre-wrap;">{{ purpose }}</pre>
        <h2>⚙️ Document Properties</h2>
        <pre style="white-space: pre-wrap;">{{ properties }}</pre>
        <h2>🐍 IronPython Scripts</h2>
        <pre style="white-space: pre-wrap;">{{ iron }}</pre>
        <h2>📜 JavaScripts</h2>
        <pre style="white-space: pre-wrap;">{{ js }}</pre>
        <h2>📊 Visualizations</h2>
        <pre style="white-space: pre-wrap;">{{ viz }}</pre>
        <h2>🗃️ Data Table Summaries (GPT-4)</h2>
        <pre style="white-space: pre-wrap;">{{ data_table_gpt_summary }}</pre>
        <h2>🪄 Filters</h2>
        <pre style="white-space: pre-wrap;">{{ filters_section }}</pre>
        <form action="{{ url_for('preview_markdown') }}" method="post">
          <button type="submit">👁️ Preview as Markdown</button>
        </form>
        <form action="{{ url_for('download_markdown') }}" method="post">
          <button type="submit">⬇️ Download as Markdown</button>
        </form>
        <br><a href="/">🔙 Upload another file</a>
        """, purpose=final_summary, properties=formatted_properties,
               iron=ironpython_summary, js=javascript_summary,
               viz=viz_summary, data_table_gpt_summary=data_table_gpt_summary,
               filters_section=filters_section)
    # GET
    return '''
    <h2>📄 Upload Dashboard DOCX File</h2>
    <form method="post" enctype="multipart/form-data">
        <input type="file" name="doc" required>
        <input type="submit" value="Generate Full Dashboard Documentation">
    </form>
    '''

@app.route("/preview", methods=["POST"])
def preview_markdown():
    markdown_doc = session.get("markdown_doc")
    if not markdown_doc:
        return "❌ No documentation to preview."
    return render_template_string("""
    <h2>Markdown Preview</h2>
    <pre style="white-space: pre-wrap;">{{ markdown_doc }}</pre>
    <form action="{{ url_for('download_markdown') }}" method="post">
      <button type="submit">⬇️ Download as Markdown</button>
    </form>
    <br><a href="/">🔙 Upload another file</a>
    """, markdown_doc=markdown_doc)

@app.route("/download", methods=["POST"])
def download_markdown():
    markdown_doc = session.get("markdown_doc")
    if not markdown_doc:
        return "❌ No documentation to download."
    return send_file(
        BytesIO(markdown_doc.encode("utf-8")),
        as_attachment=True,
        download_name="dashboard_documentation.md",
        mimetype="text/markdown"
    )

if __name__ == "__main__":
    app.run(debug=True)
